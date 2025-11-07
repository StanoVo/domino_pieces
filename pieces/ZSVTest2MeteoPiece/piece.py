from domino.base_piece import BasePiece
from .models import InputModel, OutputModel
import pandas as pd
from pathlib import Path
from docx import Document


class ZSVTest2MeteoPiece(BasePiece):
    
    def piece_function(self, input_data: InputModel):

        csv_data = []
        csv_started = False

        # Read the Word document
        input_data_file = Path(input_data.data_input_folder) / input_data.data_input_file
        doc = Document(input_data_file)
        csv_data = []
        csv_started = False
    
        # Process each paragraph in the document
        for paragraph in doc.paragraphs:
            line = paragraph.text.strip()
            # Check if we've reached the CSV data section
            if "#Data:" in line:
                csv_started = True
                continue

            # If we're in the CSV section, store the line
            if csv_started and line:
                csv_data.append(line)

        # Convert to DataFrame and write to CSV
        if csv_data:
            df = pd.DataFrame(csv_data)
            message = f"Doc with data readed successfully"
            #output_data_file = str(Path(self.results_path) / "Output_Data.csv")
            output_data_file = Path(input_data.data_output_folder) / input_data.data_output_file
            df.to_csv(output_data_file, index=False, header=False)
        #doc = Document(input_file) ??
            
        # Set display result
        self.display_result = {
            "file_type": "csv",
            "file_path": str(output_data_file)
            
        }

        # Return output
        return OutputModel(
            message=message,
            file_path=str(output_data_file)
        )